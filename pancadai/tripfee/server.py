"""
出差旅費報銷系統 - Python 伺服器
取代 GAS，使用本地 Excel + docx 範本 + docx2pdf 產出 PDF
啟動：python3 server.py --port 8080
"""

import json
import base64
import os
import re
import io
import uuid
from datetime import datetime
from pathlib import Path
from http.server import HTTPServer, BaseHTTPRequestHandler

import openpyxl
from docx import Document
from docx.shared import Pt, Cm
from docx2pdf import convert as convert_to_pdf

# ── 設定 ──
BASE_DIR = Path(__file__).parent
EXCEL_PATH = BASE_DIR / '差旅系統資料庫.xlsx'
TEMPLATE_PATH = BASE_DIR / '出差旅費報銷單_範本.docx'
OUTPUT_DIR = BASE_DIR / 'output'
OUTPUT_DIR.mkdir(exist_ok=True)


# ── Excel 操作 ──
def get_users():
    """讀取 Users 表"""
    wb = openpyxl.load_workbook(EXCEL_PATH)
    ws = wb['Users']
    users = []
    for row in ws.iter_rows(min_row=2, values_only=True):
        email, name, role, phone = row
        if email:
            users.append({
                'email': str(email).strip() if email else '',
                'name': str(name).strip() if name else '',
                'role': str(role).strip() if role else '',
                'phone': str(phone).rstrip('.0') if phone else ''
            })
    wb.close()
    return users


def append_application(data):
    """寫入 Applications 表"""
    wb = openpyxl.load_workbook(EXCEL_PATH)
    ws = wb['Applications']
    ws.append([
        data['formId'],
        datetime.now().strftime('%Y-%m-%d %H:%M:%S'),
        data.get('department', ''),
        data.get('applicant', ''),
        data.get('phone', ''),
        data.get('summary', ''),
        data.get('totalAmount', ''),
        data.get('attachments', ''),
        data.get('pdfUrl', ''),
        '待審核'
    ])
    wb.save(EXCEL_PATH)
    wb.close()


def append_details(form_id, details):
    """寫入 Details 表"""
    wb = openpyxl.load_workbook(EXCEL_PATH)
    ws = wb['Details']
    for d in details:
        ws.append([
            form_id,
            d.get('month', ''),
            d.get('day', ''),
            d.get('startLoc', ''),
            d.get('endLoc', ''),
            d.get('desc', ''),
            d.get('plane', ''),
            d.get('taxi', ''),
            d.get('train', ''),
            d.get('hotel', ''),
            d.get('meal', ''),
            d.get('other', ''),
            d.get('subtotal', '')
        ])
    wb.save(EXCEL_PATH)
    wb.close()


# ── Docx 變數取代 ──
def replace_vars_in_paragraph(para, vars_dict):
    """取代段落中的 {{變數}}"""
    for run in para.runs:
        for key, val in vars_dict.items():
            placeholder = f'{{{{{key}}}}}'
            if placeholder in run.text:
                run.text = run.text.replace(placeholder, str(val))


def fill_template(data):
    """填入範本並輸出填好的 docx"""
    doc = Document(str(TEMPLATE_PATH))

    # 主檔變數
    main_vars = {
        '申請人': data.get('applicant', ''),
        '申請單位': data.get('department', ''),
        '總金額': data.get('totalAmount', ''),
        '出差開始日期': data.get('startDate', ''),
        '出差結束日期': data.get('endDate', ''),
        '總天數': data.get('totalDays', ''),
        '用途摘要': data.get('summary', ''),
    }

    # 取代段落
    for para in doc.paragraphs:
        replace_vars_in_paragraph(para, main_vars)

    # 取代表格（含明細）
    details = data.get('details', [])
    for table in doc.tables:
        for row_idx, row in enumerate(table.rows):
            for cell in row.cells:
                for para in cell.paragraphs:
                    # 主檔變數
                    replace_vars_in_paragraph(para, main_vars)
                    # 明細變數 (n = 1..10)
                    for n in range(1, 11):
                        d = details[n - 1] if n <= len(details) else {}
                        detail_vars = {
                            f'm_{n}': d.get('month', ''),
                            f'd_{n}': d.get('day', ''),
                            f'start_{n}': d.get('startLoc', ''),
                            f'end_{n}': d.get('endLoc', ''),
                            f'desc_{n}': d.get('desc', ''),
                            f'plane_{n}': d.get('plane', ''),
                            f'taxi_{n}': d.get('taxi', ''),
                            f'train_{n}': d.get('train', ''),
                            f'hotel_{n}': d.get('hotel', ''),
                            f'meal_{n}': d.get('meal', ''),
                            f'other_{n}': d.get('other', ''),
                            f'sub_{n}': d.get('subtotal', ''),
                        }
                        replace_vars_in_paragraph(para, detail_vars)

    # 儲存填好的 docx
    form_id = data.get('formId', 'EXP-TEMP')
    filled_path = OUTPUT_DIR / f'{form_id}_filled.docx'
    doc.save(str(filled_path))
    return filled_path


def generate_pdf(docx_path, form_id):
    """將 docx 轉為 PDF 並回傳 base64，同時複製到 output 目錄"""
    pdf_path = OUTPUT_DIR / f'{form_id}.pdf'
    try:
        convert_to_pdf(str(docx_path), str(pdf_path))
    except Exception as e:
        print(f"⚠️ docx2pdf 轉換失敗: {e}，嘗試 fallback...")
        pdf_path = None

    if pdf_path and pdf_path.exists():
        with open(pdf_path, 'rb') as f:
            pdf_bytes = f.read()
        return base64.b64encode(pdf_bytes).decode('utf-8'), str(pdf_path)
    return '', ''


# ── 附件儲存 ──
def save_attachments(form_id, files):
    """儲存附件到 output 目錄"""
    att_dir = OUTPUT_DIR / f'{form_id}_attachments'
    att_dir.mkdir(exist_ok=True)
    urls = []
    for i, f in enumerate(files):
        try:
            b64_data = f['base64'].split(',')[1] if ',' in f['base64'] else f['base64']
            file_bytes = base64.b64decode(b64_data)
            fname = f['name']
            fpath = att_dir / fname
            with open(fpath, 'wb') as fout:
                fout.write(file_bytes)
            urls.append(str(fpath))
        except Exception as e:
            print(f"附件儲存失敗: {e}")
    return '\n'.join(urls)


# ── HTTP Server ──
class ReimbursementHandler(BaseHTTPRequestHandler):
    def log_message(self, format, *args):
        pass  # 靜音

    def _send_json(self, data, status=200):
        body = json.dumps(data, ensure_ascii=False).encode('utf-8')
        self.send_response(status)
        self.send_header('Content-Type', 'application/json; charset=utf-8')
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.send_header('Content-Length', str(len(body)))
        self.end_headers()
        self.wfile.write(body)

    def do_OPTIONS(self):
        self.send_response(200)
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        self.end_headers()

    def do_POST(self):
        try:
            content_length = int(self.headers.get('Content-Length', 0))
            body = self.rfile.read(content_length)
            payload = json.loads(body.decode('utf-8'))
            action = payload.get('action', '')

            print(f"[{datetime.now():%H:%M:%S}] {action} from {self.client_address}")

            if action == 'login':
                self.handle_login(payload)
            elif action == 'submit':
                self.handle_submit(payload)
            elif action == 'saveRecord':
                self.handle_save_record(payload)
            else:
                self._send_json({'success': False, 'message': f'未知動作: {action}'})

        except json.JSONDecodeError:
            self._send_json({'success': False, 'message': 'JSON 格式錯誤'}, 400)
        except Exception as e:
            print(f"❌ Error: {e}")
            self._send_json({'success': False, 'message': str(e)}, 500)

    def handle_login(self, payload):
        email = payload.get('email', '').strip()
        users = get_users()
        for u in users:
            if u['email'] == email:
                self._send_json({
                    'status': 'success',
                    'name': u['name'],
                    'role': u['role'],
                    'phone': u['phone']
                })
                return
        self._send_json({'success': False, 'message': '無存取權限'})

    def handle_submit(self, payload):
        form_id = 'EXP-' + datetime.now().strftime('%Y%m%d%H%M%S')
        payload['formId'] = form_id

        # 填入範本
        filled_docx = fill_template(payload)

        # 轉 PDF
        pdf_b64, pdf_path = generate_pdf(filled_docx, form_id)

        # 儲存附件
        att_urls = save_attachments(form_id, payload.get('files', []))

        # 寫入 Excel
        append_application({
            'formId': form_id,
            'applicant': payload.get('applicant', ''),
            'department': payload.get('department', ''),
            'phone': payload.get('phone', ''),
            'summary': payload.get('summary', ''),
            'totalAmount': payload.get('totalAmount', ''),
            'attachments': att_urls,
            'pdfUrl': pdf_path,
        })
        append_details(form_id, payload.get('details', []))

        self._send_json({
            'success': True,
            'formId': form_id,
            'base64Pdf': pdf_b64
        })

    def handle_save_record(self, payload):
        form_id = payload.get('formId', '')
        pdf_b64 = payload.get('finalPdfBase64', '')

        if pdf_b64 and form_id:
            pdf_path = OUTPUT_DIR / f'{form_id}_merged.pdf'
            try:
                pdf_bytes = base64.b64decode(pdf_b64)
                with open(pdf_path, 'wb') as f:
                    f.write(pdf_bytes)
                print(f"  💾 合併 PDF 已儲存: {pdf_path}")
            except Exception as e:
                print(f"  ⚠️ PDF 儲存失敗: {e}")

        self._send_json({'success': True})


def main():
    import argparse
    parser = argparse.ArgumentParser()
    parser.add_argument('--port', type=int, default=8080)
    args = parser.parse_args()

    server = HTTPServer(('0.0.0.0', args.port), ReimbursementHandler)
    print(f"""
╔══════════════════════════════════════╗
║  出差旅費報銷系統 - Python Server     ║
║  位址: http://127.0.0.1:{args.port}       ║
║  Excel: {EXCEL_PATH.name}            ║
║  範本: {TEMPLATE_PATH.name}  ║
╚══════════════════════════════════════╝
Ctrl+C 停止
    """)
    try:
        server.serve_forever()
    except KeyboardInterrupt:
        print("\n伺服器已停止")


if __name__ == '__main__':
    main()
