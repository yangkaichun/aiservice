#!/usr/bin/env python3
"""Build PANCREASaver 2026 AI 創新獎 提案計劃書 DOCX with cover, 4 TOCs, body, references."""

import os, sys
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement

OUT_DIR = os.path.dirname(os.path.abspath(__file__))
BLUE = "295DAA"
ORANGE = "EC7000"
WHITE = "FFFFFF"

doc = Document()

# ── Page Setup ──
for section in doc.sections:
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# ── Default style ──
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15
rPr = style.element.get_or_add_rPr()
rFonts = rPr.find(qn('w:rFonts'))
if rFonts is None:
    rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="微軟正黑體"/>')
    rPr.insert(0, rFonts)

# ── Helpers ──
def set_font(run, name='Times New Roman', ea='微軟正黑體', sz=11, bold=False, color=None):
    run.font.size = Pt(sz)
    run.bold = bold
    run.font.name = name
    rPr = run._r.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{ea}"/>')
        rPr.insert(0, rFonts)
    else:
        rFonts.set(qn('w:eastAsia'), ea)
    if color:
        run.font.color.rgb = RGBColor(*color)

def h1(doc, txt):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(10)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="8" w:space="4" w:color="{BLUE}"/></w:pBdr>'))
    run = p.add_run(txt)
    set_font(run, sz=16, bold=True, color=(0x29,0x5D,0xAA))
    return p

def h2(doc, txt):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(txt)
    set_font(run, sz=13, bold=True, color=(0x29,0x5D,0xAA))
    return p

def h3(doc, txt):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(txt)
    set_font(run, sz=11, bold=True)
    return p

def para(doc, txt, sz=11, bold=False, align=None, color=None, indent=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.space_after = Pt(6)
    if indent:
        pf.left_indent = Cm(indent)
    run = p.add_run(txt)
    set_font(run, sz=sz, bold=bold, color=color)
    return p

def add_field(p, field_code):
    """Add Word field (TOC, PAGEREF, etc.)"""
    run = p.add_run()
    r = run._r
    r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve">{field_code}</w:instrText>'))
    r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
    r2 = p.add_run('　（請於 Word 中按 Ctrl+A → F9 更新欄位）')
    set_font(r2, sz=9, color=(150,150,150))
    r3 = p.add_run()
    r3._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

def add_seq_caption(p, kind, num, caption_text):
    """SEQ field caption kind ∈ {圖,表,附件}"""
    run0 = p.add_run(f'{kind} ')
    set_font(run0, sz=10, bold=True, color=(0x29,0x5D,0xAA))
    run = p.add_run()
    r = run._r
    r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> SEQ {kind} \\* ARABIC </w:instrText>'))
    r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
    rn = p.add_run(str(num))
    set_font(rn, sz=10, bold=True, color=(0x29,0x5D,0xAA))
    re = p.add_run()
    re._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))
    rt = p.add_run(f'　{caption_text}')
    set_font(rt, sz=10)
    return p

def add_bookmark(p, name, bid):
    """Bookmark a paragraph for PAGEREF."""
    p._p.insert(0, parse_xml(f'<w:bookmarkStart {nsdecls("w")} w:id="{bid}" w:name="{name}"/>'))
    p._p.append(parse_xml(f'<w:bookmarkEnd {nsdecls("w")} w:id="{bid}"/>'))

def add_pageref(p, bookmark_name):
    """PAGEREF field - renders page number of bookmark."""
    run = p.add_run()
    r = run._r
    r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGEREF {bookmark_name} \\h </w:instrText>'))
    r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="separate"/>'))
    r2 = p.add_run('#')
    set_font(r2, sz=10)
    r3 = p.add_run()
    r3._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

def score_highlight(doc, text):
    """Orange score-highlight callout."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="FFF3E0" w:val="clear"/>'))
    pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:left w:val="single" w:sz="24" w:space="8" w:color="{ORANGE}"/></w:pBdr>'))
    run = p.add_run(f'★ 本節得分亮點　{text}')
    set_font(run, sz=10, bold=True, color=(0xEC,0x70,0x00))
    return p

def blue_highlight(doc, text):
    """Blue info banner."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    pPr = p._p.get_or_add_pPr()
    pPr.append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="E8F0FE" w:val="clear"/>'))
    run = p.add_run(f'📊 {text}')
    set_font(run, sz=10, bold=True, color=(0x29,0x5D,0xAA))
    return p

def add_table(doc, headers, rows, caption_text, table_num, bookmark_name=None, bookmark_id=None):
    """Create a styled table with SEQ caption."""
    # Caption
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_seq_caption(p, '表', table_num, caption_text)
    if bookmark_name:
        add_bookmark(p, bookmark_name, bookmark_id)
    
    ncols = len(headers)
    nrows = len(rows) + 1
    table = doc.add_table(rows=nrows, cols=ncols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Header
    for j, h in enumerate(headers):
        cell = table.cell(0, j)
        cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:shd {nsdecls("w")} w:fill="{BLUE}" w:val="clear"/>'))
        pp = cell.paragraphs[0]
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = pp.add_run(h)
        set_font(r, sz=9, bold=True, color=(255,255,255))
    
    # Body
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.cell(i+1, j)
            pp = cell.paragraphs[0]
            pp.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = pp.add_run(str(val))
            set_font(r, sz=9)
    
    doc.add_paragraph()  # spacer
    return table

def chapter_heading(doc, num, title, rubric_weight):
    """Chapter heading with rubric weight badge."""
    h1(doc, f'第{num}章　{title}')
    p = doc.add_paragraph()
    run = p.add_run(f'〔評分構面：{rubric_weight}〕')
    set_font(run, sz=10, color=(150,150,150))
    return p

# ═══════════════════════════════════════════════════
# COVER PAGE
# ═══════════════════════════════════════════════════
for _ in range(5):
    doc.add_paragraph()

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026 AI 創新獎')
set_font(run, sz=18, bold=True, color=(0x29,0x5D,0xAA))
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('2026 AI Innovation Award')
set_font(run, sz=14, color=(0x29,0x5D,0xAA))
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('仁寶賽道 · AI × 醫療照護')
set_font(run, sz=12, color=(0xEC,0x70,0x00))

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
pPr = p._p.get_or_add_pPr()
pPr.append(parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="12" w:space="1" w:color="{BLUE}"/></w:pBdr>'))

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('提 案 計 劃 書')
set_font(run, sz=24, bold=True, ea='微軟正黑體')

doc.add_paragraph()
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('PANCREASaver® 助胰見®')
set_font(run, sz=22, bold=True, ea='微軟正黑體')
run.font.color.rgb = RGBColor(0xEC, 0x70, 0x00)

p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('基於邊緣運算驅動醫療硬體之 Physical AI 胰臟癌早期偵測系統')
set_font(run, sz=13, color=(0x60,0x60,0x60))

for _ in range(3):
    doc.add_paragraph()

info_lines = [
    '參賽組別：組別一 — 臨床診斷與醫療決策創新',
    '提案單位：仲智數位健康股份有限公司（PanCAD.ai）',
    '統一編號：89183306　　成立日期：2022 年 11 月 18 日',
    '團隊代表：楊凱鈞　行銷業務總監',
    '聯絡信箱：kc.yang@pancad.ai　　手機：0967-103-155',
    '送件日期：2026 年 7 月 20 日',
]
for line in info_lines:
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(line)
    set_font(run, sz=11, ea='微軟正黑體')

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 摘要
# ═══════════════════════════════════════════════════
h1(doc, '摘　　要')

para(doc, '胰臟癌為全球最致命惡性腫瘤之一，逾八成病患確診時已屬晚期，五年存活率長年低於一成。關鍵瓶頸在於小於二公分之早期病灶於常規電腦斷層（CT）中約四成為肉眼漏診。然而，若能於腫瘤未達二公分時偵測並積極治療，五年存活率可自不足一成提升至約八成。PANCREASaver® 助胰見® 為全球首創之全自動化 Physical AI 胰臟癌 CT 輔助偵測系統，整合深度學習與影像組學雙引擎，於 CT 掃描當下即時感知微小病灶，並主動驅動終端硬體——觸發實體警示、控制 PACS 工作站優先載入高風險影像、自動產出結構化報告——實現「感知→決策→行動」之閉環。系統對小於二公分腫瘤敏感度達 92.1%，全國多中心 1,473 例驗證 AUC 達 0.95，並成功揪回 12 名放射科醫師漏診癌中 11 名。')

para(doc, '本系統已取得台灣衛福部醫療器材許可證（衛部醫器製字第007946號）及美國 FDA Breakthrough Device Designation，榮獲放射學界最高榮譽 RSNA Margulis Award（台灣首次獲獎）等九項國內外大獎，並佈局台美八件發明專利。系統已於臺大醫院、輔大醫院、博田國際醫院實際部署，累積 190 人次，通過聯新國際醫院 AI 信任中心 50 例第三方驗測。FDA 510(k) 已於 2026 年 7 月遞交申請。')

para(doc, '本提案對準 2026 AI 創新獎「仁寶賽道——臨床診斷與醫療決策創新」組別，完全契合競賽核心宗旨：「AI 是否直接影響設備行為或實體決策，而非僅分析資料」。以下各章節依競賽評分標準（場景痛點 30%／技術創新 25%／硬體整合 25%／商業模式 20%）逐項展開，以完整證據鏈證明 PANCREASaver® 已從「分析資料的被動軟體」進化為「驅動硬體行為的 Physical AI 系統」，是仁寶賽道最具獲獎實力的提案。')

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 章節目錄 (TOC field)
# ═══════════════════════════════════════════════════
h1(doc, '章節目錄')
doc.add_paragraph()
add_field(doc.add_paragraph(), r'TOC \o "1-2" \h \z \u')
# ═══════════════════════════════════════════════════
# 圖目錄（緊接章節目錄之後）
# ═══════════════════════════════════════════════════
h1(doc, '圖 目 錄')
doc.add_paragraph()
p = doc.add_paragraph()
add_seq_caption(p, '圖', 1, 'PANCREASaver 系統工作流程總覽（感知→決策→驅動閉環）')
p = doc.add_paragraph()
add_seq_caption(p, '圖', 2, '團隊組成：產學研三方鐵三角架構圖')
p = doc.add_paragraph()
add_seq_caption(p, '圖', 3, '產業場景痛點：早期胰臟癌漏診機制示意')
p = doc.add_paragraph()
add_seq_caption(p, '圖', 4, 'Physical AI 解決方案：三層感知-決策-行動閉環架構圖')
p = doc.add_paragraph()
add_seq_caption(p, '圖', 5, '系統架構：三層式邊緣運算部署圖（Edge Perception→AI Decision→HW Action）')
p = doc.add_paragraph()
add_seq_caption(p, '圖', 6, 'PANCREASaver Bridge 入口畫面與系統管理介面')
p = doc.add_paragraph()
add_seq_caption(p, '圖', 7, 'AI 輔助判讀：原始 CT 影像 vs AI 標註病灶對比')
p = doc.add_paragraph()
add_seq_caption(p, '圖', 8, '報告編輯器：結構化報告產出與醫師簽署確認')
p = doc.add_paragraph()
add_seq_caption(p, '圖', 9, '硬體平台與多中心部署地圖')
p = doc.add_paragraph()
add_seq_caption(p, '圖', 10, '服務量成長趨勢圖（2024–2026）')
p = doc.add_paragraph()
add_seq_caption(p, '圖', 11, '商業模式與市場拓展路線圖（臺灣→美國→歐盟）')

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 表目錄（緊接圖目錄之後）
# ═══════════════════════════════════════════════════
h1(doc, '表 目 錄')
doc.add_paragraph()
p = doc.add_paragraph()
add_seq_caption(p, '表', 1, '團隊核心成員與分工')
p = doc.add_paragraph()
add_seq_caption(p, '表', 2, '三大產業痛點與 Physical AI 對應解法')
p = doc.add_paragraph()
add_seq_caption(p, '表', 3, 'PANCREASaver 實際部署進度')
p = doc.add_paragraph()
add_seq_caption(p, '表', 4, '評分構面與對應章節一覽')

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 第一章　提案整體說明
# ═══════════════════════════════════════════════════
chapter_heading(doc, '一', '提案整體說明', '綜合呈現')

score_highlight(doc, '全章圍繞 Physical AI 核心主張：AI 不止分析資料，更直接驅動設備行為。以真實臨床案例破題，五項國際期刊證據鏈支撐，已取證、已落地、已救人。')

para(doc, '一位六十餘歲的女性，因持續性黃疸就診後，電腦斷層僅顯示膽管擴張，未見明確腫瘤、內視鏡超音波導引下之穿刺切片判讀為陰性。若依循常規，她極可能被歸為良性而僅作後續觀察——然而，這正是胰臟癌奪走無數生命的典型情況。經 PANCREASaver® 重新分析，系統判定為陽性並標示胰頭可疑病灶，最終手術確診，及時挽回一命。')

para(doc, '胰臟癌之所以被稱為「癌王」，在於其早期幾無症狀、病程進展迅速，逾八成病患確診時已屬晚期，五年存活率低於百分之十。電腦斷層掃描雖為主要檢測工具，然對小於二公分之早期病灶，臨床漏診率高達四成——此非個別醫療疏失，而是人眼辨識之客觀極限 [R1][R4]。然而多項研究已證實，若能於腫瘤仍小於二公分時偵測並施以積極治療，五年存活率可顯著提升至約八成 [R2]。')

para(doc, 'PANCREASaver® 助胰見® ——一套由臺灣自主研發、全球首創之 Physical AI 胰臟癌 CT 輔助偵測系統，正是針對此一醫療現場痛點而生。其結合深度學習（CNN）與影像組學（Radiomics）雙引擎，直接讀取標準 DICOM 影像，於掃描完成之際即時感知病灶，並主動驅動終端硬體——觸發閱片工作站實體警示、自動產出標記病灶影像、優先載入高風險序列於 PACS。系統對小於二公分早期腫瘤之敏感度達 92.1%，並成功偵測放射科醫師漏診病例中之 11/12 例（92%）；全國性真實世界研究（1,473 例）敏感度 89.7%、特異度 92.8%、AUC 0.95 [R1]。正因臨床價值明確，本產品獲美國 FDA 突破性醫材認定，並取得衛福部 TFDA 醫材許可（衛部醫器製字第007946號），為臺灣首張胰臟癌 AI 醫材許可證。')

para(doc, '本系統之定位並非取代醫師，而是以 AI 於臨床判讀前先行篩選，將最可疑之影像優先呈現給醫師，最終診斷仍由醫師裁定——賦予醫師一雙不受疲勞影響的輔助之眼。就流行病學而言，胰臟癌雖非發生率最高之癌別，然其高致死特性與臺灣每年約 90 萬人次之腹部 CT 檢查量，意謂大量早期病灶潛藏於既有影像中。每提升一分早期檢出率，即多挽回一名病患及一個家庭。此即本系統存在之根本價值——讓「早期發現」不再仰賴機運，而成為每一次掃描之常態。')

blue_highlight(doc, '核心數據速覽：<2cm 敏感度 92.1% ｜ 全國驗證 AUC 0.95 ｜ TFDA 許可證 ｜ FDA Breakthrough ｜ RSNA Margulis Award ｜ 台美 8 件專利 ｜ 5 篇頂尖期刊 ｜ 2 家醫學中心已上線 ｜ 累計 190 人次')

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 第二章　團隊組成與分工
# ═══════════════════════════════════════════════════
chapter_heading(doc, '二', '團隊組成與分工', '評分構面：技術創新 25%')

score_highlight(doc, '產學研鐵三角：臺大演算法 × 臺大醫院臨床 × PanCAD.ai 商轉。九項大獎、五篇頂刊、八件專利、五人核心團隊跨越醫工／數學／資工／法規／商發，五年不間斷推進——這是後進者難以複製的系統性壁壘。')

para(doc, '如同過去的醫療革命從非單一領域所能獨立成就，PANCREASaver® 之所以能自實驗室走入臨床，關鍵在於學研、臨床與商轉三方之深度整合。')

h2(doc, '技術核心：王偉仲教授（國立臺灣大學應用數學科學研究所）')
para(doc, '主導 AI 模型架構設計與高效能科學計算，為團隊五篇國際期刊之通訊作者。專精於深度學習、數值模擬與高維度資料分析，奠定 PANCREASaver 演算法之全球領先地位。其領導之 MeDA Lab 為台灣頂尖之醫學影像 AI 實驗室。')

h2(doc, '臨床驗證與需求定義：廖偉智教授（臺大醫院內科部胃腸肝膽科主治醫師暨綜合診斷部主任）')
para(doc, '以第一線消化醫學專業，確保 AI 之判讀邏輯緊扣真實臨床情境，使技術不致淪為缺乏實用性之工具。負責臨床研究設計、真實世界驗證與多專科協作，為所有五篇臨床論文之共同通訊作者。')

h2(doc, '商業化與落地推進：仲智數位健康股份有限公司（PanCAD.ai）')
para(doc, '楊凱鈞總監統籌市場拓展、KOL 導入與跨院合作。PanCAD.ai 系統工程團隊專責邊緣設備串接、RIS/PACS 底層整合與資安合規治理。團隊現為十人精實編制，橫跨醫學工程、應用數學、資訊工程、法規與商業發展等領域，形成「臺大演算法研發 → 臺大醫院臨床驗證 → PanCAD.ai 取證商轉」之產學研鐵三角。')

para(doc, '此一協作模式之成效已具體反映於其成果：累計榮獲國內外九項指標大獎、佈局台美八件發明專利，並自 2020 年起持續發表五篇同儕審查研究於國際頂尖期刊。團隊歷經全國性資料驗證、國際法規認證與臨床商轉，五年間未曾間斷推進產品落地。我們深信，此一難以複製之整合能力，正是 PANCREASaver® 能持續領先、並將學術成果實質轉化為病患福祉之根本所在。')

# Team table
add_table(doc,
    ['姓名', '角色', '所屬單位', '職稱', '核心貢獻'],
    [
        ['王偉仲', '技術總監／共同創辦人', '國立臺灣大學', '應數所教授', 'AI 模型架構、科學計算、5 篇通訊作者'],
        ['廖偉智', '臨床總監／共同創辦人', '臺大醫院', '綜合診療部主任／教授', '臨床驗證、需求定義、多專科協作'],
        ['楊凱鈞', '團隊代表／行銷業務總監', 'PanCAD.ai', '總監', '市場拓展、KOL 導入、跨院合作商轉'],
        ['PanCAD.ai 工程團隊', '系統研發與部署', 'PanCAD.ai', '10 人精實團隊', '邊緣運算、RIS/PACS 整合、資安合規'],
    ],
    '團隊核心成員與分工', 1
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 第三章　產業場景痛點
# ═══════════════════════════════════════════════════
chapter_heading(doc, '三', '產業場景痛點', '評分構面：場景痛點 30%——最高權重')

score_highlight(doc, '三大痛點直擊評分核心：（1）CT 設備純被動擷取，無主動預警——呼應「設備行為」評分；（2）放射科人力結構性過載——呼應「實務需求」；（3）早期漏診＝錯失 80% 存活率——數字直接證明問題嚴重性。每一痛點皆附文獻證據。')

para(doc, '試想：病患得以存活的關鍵資訊，其實早已存在於其 CT 影像之中；然而，僅因早期病灶小於二公分而未能為肉眼所辨識，造成錯失及早治療的契機。此即胰臟癌臨床診斷中最令人扼腕之處，亦為本提案所欲根本解決之問題：')

h2(doc, '痛點一：設備之被動性——CT 只能拍，不能警示')
para(doc, '現行 CT 設備僅具影像擷取功能，缺乏主動預警機制，無法於掃描當下提示可疑病灶；放射科醫師須於數百張切片中憑經驗逐一判讀。文獻顯示，傳統 CT 對小於二公分早期腫瘤之漏診率高達約四成，51.7% 的 CT 影像在事後回顧發現胰臟病變跡象 [R1][R4]。此非個別的醫療疏失，而是人眼辨識之客觀極限。關鍵在於：現有設備完全被動，無法在擷取影像的當下驅動任何預警行為——這正是 Physical AI 最能發揮價值的切入點。')

h2(doc, '痛點二：醫療負荷之結構性壓力——系統性困境，非經驗可解')
para(doc, '放射科醫師工作量龐大，臺灣每年約 90 萬人次腹部 CT 檢查，早期不明顯之微小病灶於高強度判讀環境下極易被忽略。此為系統性而非可仰賴經驗能克服的困境——再資深的醫師也無法保證每一份影像的每一張切片都被同等仔細地審視。')

h2(doc, '痛點三：臨床診斷之多重侷限——檢查做到極限仍可能偽陰性')
para(doc, '胰臟癌病患往往須反覆接受內視鏡超音波、穿刺切片等侵入性檢查，不僅造成身心負擔，且仍可能出現偽陰性結果，使診斷陷入僵局。本團隊已累積多例活檢陰性但 AI 正確標示病灶的臨床案例——當傳統診斷路徑走到盡頭時，AI 提供的那一個額外判斷，可能就是救命的關鍵。')

para(doc, '上述問題之代價極為沉重。臨床證據明確指出：於腫瘤小於二公分時及早偵測並積極治療，五年存活率可自個位數大幅提升至八成以上 [R2]；放射組學研究更指出 AI 可在臨床診斷前 386 天偵測到亞臨床胰臟癌 [R5]。能否及早發現，實為有效治療之分水嶺。')

para(doc, '此一痛點並非臺灣所獨有——胰臟癌於美國為第三大癌症死因且持續上升，全球放射醫學界普遍面臨人力短缺與早期漏診之共同挑戰。因此，臨床現場真正需要的是一套能於掃描當下主動偵測、即時驅動硬體警示，且不增加醫師負擔之 Physical AI 預警系統。本團隊提案所欲建立的，正是從根本改變胰臟癌篩檢典範——使被動影像設備轉為主動預警系統，讓早期發現成為制度化常態。')

# Pain point table
add_table(doc,
    ['痛點', '現狀', 'Physical AI 解法', '量化證據'],
    [
        ['設備被動性', 'CT 僅擷取影像，無預警功能', 'AI 驅動實體警示燈號／工作站聯動', '漏診率 ~40% [R1]'],
        ['人力結構性過載', '數百張切片逐張判讀', 'AI 自動篩選高風險序，優先呈現', '90 萬人次／年腹部 CT'],
        ['診斷路徑侷限', 'EUS／活檢仍可能偽陰性', 'AI 提供第三客觀意見打破僵局', '多例活檢陰性→AI 陽性→手術確診'],
    ],
    '三大產業痛點與 Physical AI 對應解法', 2
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 第四章　Physical AI 解決方案
# ═══════════════════════════════════════════════════
chapter_heading(doc, '四', 'Physical AI 解決方案', '評分構面：技術創新 25% ＋ 硬體整合 25% — 合計 50%')

score_highlight(doc, '這是全案最關鍵的章節——直接回應競賽靈魂考題：「AI 是否直接影響設備行為或實體決策」。PANCREASaver 完整實現「感知→決策→驅動」閉環，每一步都有文獻支持、有臨床案例佐證、有已部署硬體為證。')

para(doc, 'PANCREASaver® 之根本差異在於：其 AI 判讀結果可直接驅動醫院實體設備，主動介入臨床流程——而非僅止於產出一份靜態報告等待醫師有空時翻閱。此即其符合 Physical AI 定義之核心：AI 不只分析資料，更驅動設備行為 [R13]。')

h2(doc, '感知—決策—驅動 三層閉環架構')

h3(doc, '【感知層】即時感測，零人工介入')
para(doc, '系統部署於醫院端邊緣 GPU 伺服器；病患完成 CT 掃描後，DICOM 影像透過標準 C-STORE 協定即時串流進入 AI 推論管線，全自動完成窗寬窗位標準化、器官 ROI 定位與去識別化前處理。全程無須醫療人員手動圈選器官——這是臨床採用的最低門檻要求。')

h3(doc, '【決策層】雙引擎協同，精準鎖定微小病灶')
para(doc, '二維與三維卷積神經網路（改良 3D ResNet + Attention Gate）結合影像組學（Radiomics）逾百維度特徵（一階統計、形狀、GLCM、GLRLM、小波特徵），經集成模型融合判讀，於數分鐘內鎖定小於二公分之可疑病灶並輸出腫瘤位置與風險機率。雙引擎設計，專為提升對最易漏診之微小病灶的偵測敏感度——經 Lancet Digital Health 及 Radiology 等頂尖期刊驗證，對 <2cm 腫瘤敏感度達 92.1% [R1]。')

h3(doc, '【驅動層】AI 決策轉化為設備行為（Physical AI 的關鍵行動層）')
para(doc, '一旦判定為高風險，系統除產出報告外，更透過標準醫療通訊協定（HL7／FHIR／DICOM SC）與醫院 RIS 及 PACS 進行硬體與資料層級之指令連動——觸發閱片工作站發出實體警示、自動產出標記病灶影像、優先載入高風險序列於 PACS 閱片清單頂端，主動引導醫師視線。這正是 Physical AI 與一般診斷軟體的本質區別：AI 的決策直接轉化為設備的物理行為。')

h2(doc, '臨床案例實證')

para(doc, '案例一：一名七十餘歲女性，新發糖尿病併 CA 19-9 輕度上升，CT 肉眼未見明確腫瘤，經 PANCREASaver® 判定陽性並標示胰頸可疑區域，術後病理確診為 1.8 公分早期胰臟腺癌（PDAC）。此類個案於傳統流程中極易被漏診而錯失手術良機。')

para(doc, '案例二：一名六十餘歲女性，出現進行性黃疸、CT 顯示膽管擴張而無明顯胰腺腫瘤，內視鏡超音波及活檢均無診斷性結論，但透過 PANCREASaver® 診斷為陽性並顯示胰頭可疑區域，最終手術切除後經病理證實為 PDAC。此案例特別凸顯了 AI 在傳統診斷路徑走到盡頭時的獨特價值。')

para(doc, '綜言之，本系統將「被動之診斷軟體」轉化為「由 AI 主動驅動醫療硬體之 Physical AI」，於不改變醫師既有閱片習慣之前提下，實現決策驅動實體行為之變革——使每一次 CT 掃描皆成為主動攔截早期病灶的契機。')

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 第五章　具體任務與複審 Demo 規劃
# ═══════════════════════════════════════════════════
chapter_heading(doc, '五', '具體任務與複審 Demo 規劃', '評分構面：硬體整合 25%（複審現場 DEMO）')

score_highlight(doc, '六項可 Demo 任務全部對準複審評分——含「環境干擾與隨機變數測試」準備。現場實機運行，不是放影片。已備妥 Siemens／GE 雙機型案例及偽陽性挑戰陰性案例。')

para(doc, '為驗證系統之實際效能，本團隊將於複審階段進行現場實機示範，展示 AI 判讀如何驅動實體硬體反應。具體可示範之六大任務如下：')

tasks = [
    ('🚩 任務一：即時硬體警示聯動', '邊緣伺服器接收 CT 影像後，於數分鐘內辨識 <2cm 病灶，成功觸發實體警示裝置（警示燈號／工作站聯動），驗證 AI 決策可直接驅動硬體。'),
    ('🚩 任務二：自動報告產出與系統驅動', '系統自動產出結構化檢測報告並寫入 RIS，同步傳輸至 PACS，全程無須人為介入。'),
    ('🚩 任務三：病灶可視化標示', '系統並列呈現原始影像與 AI 標註影像，以高對比視覺標示腫瘤位置，除引導醫師判讀外，亦有助於醫病溝通。'),
    ('🚩 任務四：臨床流程無縫整合', '於完全不改變醫師既有閱片習慣之前提下，驗證系統即插即用之特性——此為醫師長期採用之關鍵。'),
    ('🚩 任務五：真實案例重現', '以去識別化之真實個案示範：1.8 公分早期 PDAC 及活檢偽陰性 1 公分病灶，展示 AI 如何攔截傳統檢查難以察覺之早期病變。'),
    ('🚩 任務六：效能實測與隨機變數測試', '現場展示單份逾五百張切片之 CT 於五分鐘內完成推論。另備不同 CT 機型（Siemens、GE 各一例）及含偽陽性挑戰之陰性案例，完整回應競賽「環境干擾與隨機變數測試」要求。'),
]
for title, desc in tasks:
    h3(doc, title)
    para(doc, desc)

para(doc, '上述任務皆非實驗室假設，而是系統於臺大醫院、輔大醫院實際運作之常態：截至 2026 年 6 月，累計臨床輔助判讀達 190 人次（臺大 35＋輔大 155），博田國際醫院預約累積 11 例，聯新國際醫院透過衛福部「負責任 AI 執行中心」完成 50 例第三方驗測。我們能以真實個案重現完整臨床流程——影像接收、判讀、警示觸發與報告產出，每一環節皆為導入院所的例行運作。')

para(doc, '我們深信：令評審親眼見證 AI 偵測病灶、工作站觸發警示、報告自動產出之完整歷程，遠較任何簡報論述更具說服力。這一刻的親眼見證，勝過千言萬語。')

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 第六章　系統架構與 AI 整合設計
# ═══════════════════════════════════════════════════
chapter_heading(doc, '六', '系統架構與 AI 整合設計', '評分構面：技術創新 25% ＋ 硬體整合 25%')

score_highlight(doc, '三層架構（邊緣感知→AI 決策→醫療系統驅動）完整呈現 Physical AI 的工程實現。PANCREASaver Bridge 中介軟體為自研核心。S-SDLC 安全開發、去識別化、LDAP/AD 整合——資安合規一次到位。')

para(doc, '醫療現場對系統之要求，在於判讀精準、反應即時，且能安全地於院內邊緣運算環境穩定運行。為此，本系統採全程隱身於醫院封閉內網之設計，不依賴外部網路，確保最高等級之資訊安全。')

h2(doc, '三層式邊緣運算架構')

h3(doc, '一、邊緣感知層（Perception）')
para(doc, 'CT 產出 DICOM 影像後，透過標準 DICOM C-STORE 協定直連院內本地 GPU 伺服器（NVIDIA RTX 4090／T4），全自動完成窗寬窗位標準化、器官 ROI 定位與去識別化。感測對象為標準化 CT 影像序列，無須醫療人員手動介入。亦支援雲端部署方案供未來擴充。')

h3(doc, '二、AI 決策層（Decision）')
para(doc, '二維與三維卷積神經網路（改良 3D ResNet + Attention Gate）結合影像組學逾百維度特徵，經集成模型融合判讀。模型於全國性真實世界資料完成驗證：敏感度 89.7%、特異度 92.8%、AUC 0.95 [R1]。以真實世界資料——而非實驗室精選資料——進行驗證，對臨床信任至關重要。')

h3(doc, '三、醫療系統驅動層（Action）')
para(doc, '判定陽性後，透過標準醫療通訊協定（HL7／FHIR／DICOM SC）與內網交換指令，同步完成：報告產出 → RIS 寫入 → PACS 畫面整合與警示。PANCREASaver Bridge 為自研核心中介軟體，負責各層間的訊息路由、事件日誌與異常重試機制。Report Editor 前端（Vue.js）提供醫師審閱與簽署介面。')

h2(doc, '資訊安全與合規治理')
para(doc, '系統導入安全軟體開發生命週期（S-SDLC）；病患影像於推論前完成端點去識別化；遵循《個人資料保護法》及衛福部負責任 AI 治理框架；使用者身分驗證透過 LDAP／AD 整合院內帳號系統並具完整操作稽核日誌。')

h2(doc, '技術壁壘')
para(doc, '以台美八件發明專利佈局（台灣 6 件＋美國 4 件），採系統級與方法級雙重權利範圍撰寫，確認於臺灣、美國、歐盟及中國均無侵權風險。整體架構之核心優勢在於：完全融入放射科醫師既有工作流程——易於介接、推論迅速、資料安全。此三項條件之達成，方為醫療 AI 得以進入臨床並長期留存之前提。')

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 第七章　硬體平台與實體驗證規劃
# ═══════════════════════════════════════════════════
chapter_heading(doc, '七', '硬體平台與實體驗證規劃', '評分構面：硬體整合 25%——關鍵得分章節')

score_highlight(doc, '已部署、已上線、已驗證——不是「規劃中」或「預計」。臺大＋輔大＋博田三家醫院實際運行，190 人次，聯新 AI 信任中心 50 例第三方驗測，FDA 510(k) 已遞交。用真實世界的運行數據回應「實體驗證」評分。')

para(doc, '任何無法於真實臨床環境落地之技術，終究僅為實驗室之理論。本團隊之驗證策略，係將邊緣運算硬體直接部署於臨床第一線，於真實病患與實際流程中接受檢驗。')

h2(doc, '硬體平台規格')
hw_specs = [
    '高效能 GPU 邊緣運算主機（NVIDIA RTX 4090／T4 等級），單份逾五百張切片 CT 可於五分鐘內完成推論',
    '內建實體警示模組（警示燈號／工作站 API 聯動）',
    '標準醫療網路介接協議（DICOM／HL7／FHIR），與院內既有 PACS 及 RIS 無縫整合',
    '作業系統：Ubuntu 22.04 LTS，支援雙 GPU failover 備援',
    'PANCREASaver Bridge 中介軟體（自研）：訊息路由、事件日誌、異常重試機制',
]
for spec in hw_specs:
    para(doc, f'• {spec}')

h2(doc, '實體驗證四階段')

para(doc, '第一階段｜影子測試（已完成）：團隊已於頂尖醫學中心之後台同步運行系統，在完全不干擾既有醫療流程之前提下，驗證 AI 驅動硬體警示之速度、準確度與穩定度。')

para(doc, '第二階段｜真實世界營運驗證（已完成並持續進行）：系統已於臺大醫院總院、輔大醫院聖路加健檢中心正式上線，截至 2026 年 6 月累計臨床輔助判讀 190 人次（臺大 35＋輔大 155）；博田國際醫院 2026 年 6 月起導入，預約累積 11 例；聯新國際醫院透過衛福部「負責任 AI 執行中心」進行 50 例第三方驗測；彰化基督教醫院議價中。此等由院方主動發起之導入與驗測，印證硬體落地之穩定性與臨床價值。')

para(doc, '第三階段｜持續監控與前瞻性研究（進行中）：臺大研發團隊持續監測模型於真實硬體環境下之靈敏度；規劃 2026 Q3–2027 Q2 於 5 家以上區域醫院進行 500+ 例前瞻性臨床研究。')

para(doc, '第四階段｜法規驗證（已完成＋進行中）：系統已通過 TFDA 醫材查驗登記（衛部醫器製字第007946號），為臺灣首張胰臟癌 AI 醫材許可證；美國 FDA 510(k) 已於 2026 年 7 月遞交申請，以國際最高標準再次驗證軟硬體整合之安全性與有效性。')

para(doc, '將技術直接推向真實臨床戰場，是最誠實亦最嚴苛之驗證方式。本團隊選擇此途，正因唯有如此，方能對得起每一位將生命託付於醫療現場之病患。已上線、已實際輔助臨床、已推進國際認證——此即本團隊對「落地」二字最務實之交代。')

# Deployment table
add_table(doc,
    ['部署機構', '狀態', '累積案例', '備註'],
    [
        ['臺大醫院總院', '✅ 正式上線', '35 人次', '放射科即時 AI 輔助判讀'],
        ['輔大醫院聖路加健檢中心', '✅ 正式上線', '155 人次', '臺灣最大腹部 MRI 健檢中心'],
        ['博田國際醫院', '✅ 2026.06 導入', '11 例（預約）', '南部市場擴張首站'],
        ['聯新國際醫院', '🔄 驗測中', '50 例', '衛福部 AI 信任中心第三方驗測'],
        ['彰化基督教醫院', '📋 議價中', '—', '中部市場拓展目標'],
    ],
    'PANCREASaver 實際部署進度', 3, 'tbl_deploy', 100
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 第八章　商業模式與市場發展規劃
# ═══════════════════════════════════════════════════
chapter_heading(doc, '八', '商業模式與市場發展規劃', '評分構面：商業模式 20%')

score_highlight(doc, 'SaaS 年約訂閱制＋仁寶 OEM 預載雙軌策略。臺灣 90 萬人次／年腹 CT × 全球 130 億美元市場。已取得 Term Sheet、建置 QMS、聯新加速器資源。後進者壁壘：率先研發→率先臨床→率先取證。')

para(doc, '本產品之市場拓展已具實質進展，而非停留於規劃階段。')

h2(doc, '導入實績')
para(doc, '系統已導入臺大醫院、輔大醫院聖路加健檢中心；博田國際醫院為南部擴張首站（2026 年 6 月起）；聯新國際醫院透過衛福部「負責任 AI 執行中心」進行 50 例第三方驗測；彰化基督教醫院議價中。截至 2026 年 6 月累計臨床輔助判讀 190 人次。190 人次——誠實地說，不多。但胰臟癌不是常見癌別，每一次判讀都可能攔截一條命，我們看的不是絕對量，是成長斜率：臺大 35＋輔大 155＋博田 11 預約，逐月攀升。')

h2(doc, '商業模式：雙軌策略')
para(doc, '【SaaS 年約訂閱制】依 CT 掃描量分級計費。醫院導入後可新增針對高致死率癌症之自費篩檢項目，於不增加放射科醫師工作量之前提下創造新收益。早期發現可避免晚期龐大之化療、標靶與住院支出，實質為病患家庭與國家健保節省成本。')
para(doc, '【仁寶 OEM 預載策略】將 PANCREASaver 預先部署於仁寶 Edge AI 伺服器或醫規工作站，以「硬體＋AI」套裝方案銷售，降低醫院導入門檻，創造企業賽道協同效益——此為本提案與仁寶賽道之獨特契合點。')

h2(doc, '市場規模')
para(doc, '臺灣每年約 90 萬人次接受腹部 CT 檢查；全球 AI 醫學影像市場預估於 2030 年達 130 億美元，年複合成長率逾 35%。PANCREASaver 之 FDA 510(k) 已遞交，規劃 2027 年啟動美國初步商業化、2028 年歐盟 CE MDR 申請。後續研發將推進無顯影劑 CT（NC-CT）與磁振造影（MRI）之新模組，擴大適應症覆蓋。')

h2(doc, '公司永續與競爭壁壘')
para(doc, '現有十人精實團隊、已取得投資意向書（Term Sheet），並正建置自主醫療器材品質管理系統（QMS），預計 2026 Q4–2027 H1 取證，屆時將具備完整醫材製造能力；另獲 2026 聯新數位健康加速器資源挹注。')
para(doc, '於胰臟癌 AI 此一領域，率先完成研發、率先進入臨床、率先取得國際認證者，將建立後進者難以跨越之門檻。本團隊已領先數年——兩家醫學中心導入、190 人次實證、九項國內外大獎、八件發明專利及五篇國際頂尖期刊——均非後進者短期所能補足。這不只是一具商業爆發力的產品，更是一項足以挽救眾多生命的志業。')

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 對照評分構面總結
# ═══════════════════════════════════════════════════
h1(doc, '評分構面對照總表')

para(doc, '以下對照表呈現本提案各章節與競賽四大評分構面之對應關係，供評審快速查閱：')

add_table(doc,
    ['評分構面', '權重', '對應章節', 'PANCREASaver 得分關鍵'],
    [
        ['場景痛點與解決方案', '30%', '第一章、第三章', '三層痛點（設備被動／人力過載／診斷侷限）→ Physical AI 主動驅動解法；<2cm 敏感度 92.1%；AUC 0.95'],
        ['技術創新與系統架構', '25%', '第四章、第六章', '感知→決策→驅動閉環；CNN＋Radiomics 雙引擎；5 篇頂刊驗證；台美 8 件專利佈局'],
        ['硬體整合與實作計畫', '25%', '第五章、第七章', '已部署 3 家醫院實際運行 190 人次；6 項可 Demo 任務含隨機變數測試；FDA 510(k) 已遞交'],
        ['商業模式與市場性', '20%', '第八章', 'SaaS＋OEM 雙軌；仁寶預載策略；全球 130 億美元市場；Term Sheet＋QMS 建置中'],
    ],
    '評分構面與對應章節一覽', 4, 'tbl_scoring', 101
)

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 參考文獻
# ═══════════════════════════════════════════════════
h1(doc, '參考文獻')

refs = [
    ('[R1] ★', 'Chen PT, Chang D, Wu T, et al. "Pancreatic Cancer Detection on CT Scans with Deep Learning: A Nationwide, Real-World Study." Radiology, 2023; 308(2): e220152. DOI: 10.1148/radiol.220152. PMID: 36926666.',
     'PANCREASaver 核心臨床驗證文獻。'),
    ('[R2] ★', 'Egawa S, Takeda K, et al. "Clinicopathological aspects of small pancreatic cancer." Pancreas, 2004; 28(3): 235-240. DOI: 10.1097/00006676-200404000-00004.',
     '<2cm 胰臟癌積極治療，五年存活率可達 58–80%。'),
    ('[R3]', 'Sung H, Ferlay J, et al. "Global Cancer Statistics 2020: GLOBOCAN." CA: A Cancer Journal for Clinicians, 2021; 71(3): 209-249. DOI: 10.3322/caac.21660.',
     '全球癌症統計。'),
    ('[R4] ★', 'Hoogenboom SA, et al. "Prevalence, features, and explanations of missed and misinterpreted pancreatic cancer on imaging." Abdominal Radiology, 2022; 47(12): 4160–4172. DOI: 10.1007/s00261-022-03671-6.',
     '51.7% CT 事後發現胰臟病變跡象被遺漏。'),
    ('[R5] ★', 'Mukherjee S, et al. "Radiomics-based ML Models Can Detect Pancreatic Cancer on Prediagnostic CT at a Substantial Lead Time." Gastroenterology, 2022; 163(5): 1234-1245. DOI: 10.1053/j.gastro.2022.07.015.',
     'AI 可在臨床診斷前 386 天偵測亞臨床胰臟癌。'),
    ('[R6]', 'Cao K, et al. "Large-scale pancreatic cancer detection via non-contrast CT and deep learning." Nature Medicine, 2023; 29: 3033–3043. DOI: 10.1038/s41591-023-02640-w.',
     '非對照 CT 大規模偵測胰臟癌可行性。'),
    ('[R7]', 'Liu KL, Wu TH, Chen PT, et al. "Deep learning to distinguish pancreatic cancer from benign pancreatic conditions on CT." Radiology: Imaging Cancer, 2020; 2(4): e190096. DOI: 10.1148/rycan.2020190096.',
     'PANCREASaver 前期方法論。'),
    ('[R8]', 'Liao WC, Chen WY, et al. "AI detects missed pancreatic cancer at abdominal CT." The Lancet Digital Health, 2021; 3(6): e346-e347. DOI: 10.1016/S2589-7500(21)00060-X.',
     'AI 回溯性揪出漏診胰臟癌。'),
    ('[R9]', 'Chen WY, Liao WC, et al. "Radiomics-based ML for differentiating pancreatic cancer from mass-forming chronic pancreatitis on CT." BMC Cancer, 2023; 23: 345. DOI: 10.1186/s12885-023-10777-5.',
     '放射組學於胰臟癌鑑別診斷。'),
    ('[R10]', 'Yao L, et al. "A review of deep learning and radiomics approaches for pancreatic cancer diagnosis." Current Opinion in Gastroenterology, 2023; 39(5): 436–447. DOI: 10.1097/MOG.0000000000000966.',
     '深度學習與放射組學於胰臟癌診斷綜述。'),
    ('[R11]', 'Lopez-Ramirez F, et al. "Early detection of pancreatic cancer on CT: advancements with deep learning." Radiology Advances, 2025; 2(5): umaf028. DOI: 10.1093/radadv/umaf028.',
     '深度學習於胰臟癌早期偵測最新進展。'),
    ('[R12]', 'Pongprasobchai S, et al. "Long-term survival and prognostic indicators in small pancreatic cancer." Pancreatology, 2008; 8(6): 587-592. DOI: 10.1159/000161009.',
     '<2cm 胰臟癌長期存活與預後指標。'),
    ('[R13]', 'Ferrone CR, et al. "Pancreatic adenocarcinoma: the actual 5-year survivors." J Gastrointest Surg, 2008; 12(4): 701-706. DOI: 10.1007/s11605-007-0384-8.',
     '胰臟癌實際五年存活者分析。'),
]
for code, citation, note in refs:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(f'{code} {citation}')
    set_font(run, sz=9)
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_after = Pt(6)
    p2.paragraph_format.left_indent = Cm(1)
    run2 = p2.add_run(f'※ {note}')
    set_font(run2, sz=9, color=(100,100,100))

doc.add_page_break()

# ═══════════════════════════════════════════════════
# 附件目錄
# ═══════════════════════════════════════════════════
h1(doc, '附件目錄')
doc.add_paragraph()

attachments = [
    ('附件一', 'PANCREASaver TFDA 醫療器材許可證（衛部醫器製字第007946號）——臺灣首張胰臟癌 AI 醫材許可證', 'att_tfda', 200),
    ('附件二', '美國 FDA Breakthrough Device Designation 認證函', 'att_fda', 201),
    ('附件三', 'RSNA Margulis Award 及歷年國內外獲獎證明（共 9 項）', 'att_awards', 202),
    ('附件四', '學術期刊論文（Radiology／Lancet DH／BMC Cancer／JGH／Radiology: Imaging Cancer 共 5 篇）', 'att_papers', 203),
    ('附件五', '智財權組合（台灣 6 件＋美國 4 件專利文件）', 'att_patents', 204),
    ('附件六', '臨床輔助判讀完整報告範例（陽性／陰性各 25 頁）', 'att_reports', 205),
    ('附件七', '2026 AI 創新獎參賽同意書（已用印）', 'att_consent', 206),
    ('附件八', '提案簡報投影片（Pitch Deck，11 頁）', 'att_deck', 207),
]

for kind, desc, bm_name, bm_id in attachments:
    p = doc.add_paragraph()
    add_seq_caption(p, '附件', attachments.index((kind, desc, bm_name, bm_id)) + 1, desc)
    add_bookmark(p, bm_name, bm_id)

# ── SAVE ──
output = os.path.join(OUT_DIR, 'PANCREASaver_2026AI創新獎_提案計劃書_完整版.docx')
doc.save(output)
print(f'✅ 已產出：{output}')
print(f'📏 檔案大小：{os.path.getsize(output):,} bytes')
print(f'📋 後續步驟：在 Microsoft Word 中開啟 → Ctrl+A → F9 更新所有欄位 → 另存為 PDF')
